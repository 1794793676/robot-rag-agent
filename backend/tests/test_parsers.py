from docx import Document
from openpyxl import Workbook
import pytest
import xlwt

from app.rag.parsers import ParsedSection, parse_document


def test_docx_preserves_block_order_and_heading_context(tmp_path):
    path = tmp_path / "safety.docx"
    document = Document()
    document.add_heading("安全规范", level=1)
    document.add_paragraph("操作前关闭主电源。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "电源"
    table.cell(1, 1).text = "关闭"
    document.add_paragraph("确认指示灯熄灭。")
    document.save(path)

    sections = parse_document(path, "docx")

    assert sections == [
        ParsedSection(text="操作前关闭主电源。", heading="安全规范"),
        ParsedSection(
            text="| 项目 | 要求 |\n| --- | --- |\n| 电源 | 关闭 |",
            heading="安全规范",
        ),
        ParsedSection(text="确认指示灯熄灭。", heading="安全规范"),
    ]


def test_docx_preserves_nested_heading_path(tmp_path):
    path = tmp_path / "nested-heading.docx"
    document = Document()
    document.add_heading("一级", level=1)
    document.add_paragraph("一级正文。")
    document.add_heading("二级", level=2)
    document.add_paragraph("二级正文。")
    document.save(path)

    sections = parse_document(path, "docx")

    assert sections == [
        ParsedSection(text="一级正文。", heading="一级"),
        ParsedSection(text="二级正文。", heading="一级 / 二级"),
    ]


def test_docx_recognizes_short_chinese_numbered_normal_paragraph_as_heading(
    tmp_path,
):
    path = tmp_path / "numbered-heading.docx"
    document = Document()
    document.add_paragraph("一、概念")
    document.add_paragraph("这是正文。")
    document.save(path)

    sections = parse_document(path, "docx")

    assert sections == [ParsedSection(text="这是正文。", heading="一、概念")]


def test_docx_keeps_punctuated_chinese_numbered_steps_as_body_text(tmp_path):
    path = tmp_path / "numbered-steps.docx"
    document = Document()
    document.add_paragraph("一、关闭电源。")
    document.add_paragraph("二、确认指示灯熄灭。")
    document.save(path)

    sections = parse_document(path, "docx")

    assert sections == [
        ParsedSection(text="一、关闭电源。"),
        ParsedSection(text="二、确认指示灯熄灭。"),
    ]


def test_xlsx_preserves_sheet_name_headers_and_rows(tmp_path):
    path = tmp_path / "robot-params.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "参数表"
    sheet.append(["部件", "参数", "数值"])
    sheet.append(["电池", "额定电压", "48V"])
    sheet.append(["电机", "最大转速", "3000rpm"])
    empty_sheet = workbook.create_sheet("空表")
    empty_sheet.append([None, None])
    workbook.save(path)

    sections = parse_document(path, "xlsx")

    assert sections == [
        ParsedSection(
            text=(
                "| 部件 | 参数 | 数值 |\n"
                "| --- | --- | --- |\n"
                "| 电池 | 额定电压 | 48V |\n"
                "| 电机 | 最大转速 | 3000rpm |"
            ),
            heading="工作表 参数表 / 行 1-3",
        )
    ]


def test_xls_preserves_sheet_name_headers_and_rows(tmp_path):
    path = tmp_path / "legacy-params.xls"
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("参数表")
    rows = [
        ["部件", "参数", "数值"],
        ["电池", "额定电压", "48V"],
        ["电机", "最大转速", "3000rpm"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            sheet.write(row_index, column_index, value)
    workbook.save(path)

    sections = parse_document(path, "xls")

    assert sections == [
        ParsedSection(
            text=(
                "| 部件 | 参数 | 数值 |\n"
                "| --- | --- | --- |\n"
                "| 电池 | 额定电压 | 48V |\n"
                "| 电机 | 最大转速 | 3000rpm |"
            ),
            heading="工作表 参数表 / 行 1-3",
        )
    ]


def test_pdf_without_text_layer_reports_no_ocr_support(tmp_path, monkeypatch):
    from app.rag import parsers

    class EmptyPage:
        def extract_text(self):
            return ""

    class EmptyPdfReader:
        def __init__(self, path):
            self.pages = [EmptyPage()]

    monkeypatch.setattr(parsers, "PdfReader", EmptyPdfReader)
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    with pytest.raises(parsers.DocumentParseError) as exc_info:
        parsers.parse_document(pdf_path, "pdf")

    assert str(exc_info.value) == "当前版本不支持 OCR，请上传可复制文本的 PDF、docx 或 txt 文件。"
